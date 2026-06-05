"""
export_word.py
Render the ranked shortlist as a polished Word document (.docx) for the
investment team. Follows skills/06_brief_document_standard.md and pulls ALL
styling from config/report_style.json (single source of truth).

House style: classic serif (Georgia), restrained monochrome palette, a real
bordered summary table, accent rule dividers, and a confidential footer with
page numbers. All text is plain ASCII (no smart quotes / em-dashes).
"""

import json
import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from generate_report import (build_summary, _template_rationale,
                             _clean_domain, clean_text)


def _load_style(path="config/report_style.json"):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}


def _hex(c):
    return RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))


def _font(run, font, size, bold=False, color=None, italic=False):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = _hex(color)


def _bottom_border(paragraph, color, size=6):
    """Add a horizontal rule under a paragraph (used as a divider)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def _shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def _add_hyperlink(paragraph, url, text, color):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyper = OxmlElement("w:hyperlink")
    hyper.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), color); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    hyper.append(run)
    paragraph._p.append(hyper)


def _footer(section, text, font, color):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text + "      Page ")
    _font(r, font, 8, color=color)
    fld1 = OxmlElement("w:fldSimple"); fld1.set(qn("w:instr"), "PAGE")
    run_xml = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    co = OxmlElement("w:color"); co.set(qn("w:val"), color); rpr.append(co)
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "16"); rpr.append(sz)
    run_xml.append(rpr); t = OxmlElement("w:t"); t.text = "1"; run_xml.append(t)
    fld1.append(run_xml); p._p.append(fld1)


def export_to_word(ranked, config, output_dir, input_name=""):
    style = _load_style()
    font = style.get("font", "Georgia")
    body = style.get("body_size_pt", 11)
    col = style.get("colors", {})
    primary = col.get("primary", "1A1A1A")
    accent = col.get("accent", "8A7357")
    muted = col.get("muted", "595959")
    faint = col.get("faint", "8C8C8C")

    doc = Document()
    doc.styles["Normal"].font.name = font
    doc.styles["Normal"].font.size = Pt(body)

    sec = doc.sections[0]
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1); sec.right_margin = Inches(1)
    _footer(sec, clean_text(style.get("footer_text", "Confidential")), font, muted)

    p = doc.add_paragraph()
    _font(p.add_run(clean_text(style.get("title", "Weekly Sourcing Brief"))),
          font, 26, bold=True, color=primary)
    if style.get("subtitle"):
        p = doc.add_paragraph()
        _font(p.add_run(clean_text(style["subtitle"]).upper()), font, 12, color=accent)
    p = doc.add_paragraph()
    src = f"  Source: {input_name}." if input_name else ""
    _font(p.add_run(f"Top {len(ranked)} qualified companies "
                    f"(minimum score {config['min_score_threshold']}).{src}"), font, 10, color=muted)
    _bottom_border(p, accent, size=12)
    doc.add_paragraph("")

    if style.get("show_summary_section", True):
        s = build_summary(ranked, config, style)
        hp = doc.add_paragraph()
        _font(hp.add_run("Portfolio Summary"), font, 15, bold=True, color=primary)

        rows = [("Tiers", s["tiers_str"]), ("Sector concentration", s["sectors_str"]),
                ("Stages", s["stages_str"])]
        table = doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.5)
        for ri, (k, v) in enumerate(rows):
            c0, c1 = table.rows[ri].cells
            c0.width = Inches(2.0); c1.width = Inches(4.5)
            _shade(c0, "F2EFEA")
            c0.paragraphs[0].clear()
            _font(c0.paragraphs[0].add_run(k), font, 10, bold=True, color=primary)
            c1.paragraphs[0].clear()
            _font(c1.paragraphs[0].add_run(v), font, 10, color="000000")

        pp = doc.add_paragraph()
        _font(pp.add_run("Top picks:  "), font, body, bold=True, color=primary)
        _font(pp.add_run(s["picks_str"]), font, body, color="000000")
        doc.add_paragraph("")

    for i, row in ranked.iterrows():
        tier = clean_text(row.get("qualification_tier", ""))
        name = clean_text(row.get("Company Name", ""))
        hp = doc.add_paragraph()
        _font(hp.add_run(f"{i+1}.  {name}"), font, 15, bold=True, color=primary)
        _font(hp.add_run(f"      Score {int(row['total_score'])}"), font, 12, bold=True, color=accent)
        if tier:
            _font(hp.add_run(f"   ({tier})"), font, 10, color=muted)

        meta = "  |  ".join(x for x in [clean_text(row.get("Growth Stage", "")),
                                         clean_text(row.get("Industry", "")),
                                         clean_text(row.get("HQ Location", ""))]
                            if x and x != "nan")
        if meta:
            p = doc.add_paragraph()
            _font(p.add_run(meta), font, 9, color=faint, italic=True)

        dom = _clean_domain(row)
        if dom:
            p = doc.add_paragraph()
            _font(p.add_run("Website:  "), font, body, bold=True, color=primary)
            _add_hyperlink(p, f"https://{dom}", dom, accent)

        p = doc.add_paragraph()
        _font(p.add_run(_template_rationale(row)), font, body, color="000000")

        if "outreach_action" in row:
            p = doc.add_paragraph()
            _font(p.add_run("Action:  "), font, body, bold=True, color=primary)
            _font(p.add_run(f"{clean_text(row['outreach_action'])}. {clean_text(row.get('timing_note', ''))}"),
                  font, body, color="000000")
            steps = clean_text(row.get("diligence_steps", "")).split(" | ")
            if steps and steps[0]:
                p = doc.add_paragraph()
                _font(p.add_run("Diligence:"), font, body, bold=True, color=primary)
                for st in steps:
                    bp = doc.add_paragraph(style="List Bullet")
                    _font(bp.add_run(st), font, body, color="000000")

        div = doc.add_paragraph()
        _bottom_border(div, "DDDDDD", size=4)

    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, "investor_brief.docx")
    doc.save(out_path)
    return out_path
